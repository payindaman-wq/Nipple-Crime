"""
Nipple Crime SOP Generator
Creates a formatted DOCX SOP matching the master template (Tr3).
Usage: python scripts/create_sop.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree
import os

NS_WP  = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
NS_A   = 'http://schemas.openxmlformats.org/drawingml/2006/main'
NS_PIC = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
NS_R   = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_run(para, text, bold=False, size=None, color=None, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_floating_image(paragraph, image_path, width_emu, height_emu, pos_h_emu, pos_v_emu):
    """
    Add a floating (anchored) image to a paragraph by:
    1. Adding it inline to register the relationship and get the rId
    2. Replacing the inline XML with an anchor at the specified position
    """
    run = paragraph.add_run()
    run.add_picture(image_path, width=Emu(width_emu), height=Emu(height_emu))

    w_drawing = run._element.find(qn('w:drawing'))
    inline = w_drawing.find('{%s}inline' % NS_WP)

    # Extract rId from blip element
    blip = inline.find('.//{%s}blip' % NS_A)
    r_id = blip.get('{%s}embed' % NS_R)

    anchor_xml = (
        '<wp:anchor xmlns:wp="{wp}" xmlns:a="{a}" xmlns:pic="{pic}" xmlns:r="{r}" '
        'distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="3" '
        'behindDoc="0" locked="0" layoutInCell="0" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="column"><wp:posOffset>{ph}</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="paragraph"><wp:posOffset>{pv}</wp:posOffset></wp:positionV>'
        '<wp:extent cx="{cx}" cy="{cy}"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapSquare wrapText="largest"/>'
        '<wp:docPr id="2" name="BMlogo"/>'
        '<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        '<pic:pic><pic:nvPicPr>'
        '<pic:cNvPr id="2" name="BMlogo"/>'
        '<pic:cNvPicPr><a:picLocks noChangeAspect="1" noChangeArrowheads="1"/></pic:cNvPicPr>'
        '</pic:nvPicPr>'
        '<pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        '<pic:spPr bwMode="auto">'
        '<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom><a:noFill/>'
        '</pic:spPr>'
        '</pic:pic></a:graphicData></a:graphic>'
        '</wp:anchor>'
    ).format(
        wp=NS_WP, a=NS_A, pic=NS_PIC, r=NS_R,
        ph=pos_h_emu, pv=pos_v_emu,
        cx=width_emu, cy=height_emu,
        rid=r_id,
    )

    anchor = etree.fromstring(anchor_xml)
    w_drawing.remove(inline)
    w_drawing.append(anchor)


def create_sop(
    output_path,
    sop_number,
    sop_title,
    department,
    version,
    effective_date,
    last_updated,
    sections,
    nc_logo_path="Images/NC logo.png",
    bm_logo_path="Images/BM logo.jpg",
):
    doc = Document()

    page_section = doc.sections[0]
    page_section.top_margin    = Inches(0.5)
    page_section.bottom_margin = Inches(0.75)
    page_section.left_margin   = Inches(1)
    page_section.right_margin  = Inches(1)

    # =========================================================
    # HEADER: NC logo inline left + BM logo floating right
    # Dimensions and positions match the Tr3 master template exactly.
    #   NC logo  — inline:  ~3.16" wide x 0.90" tall
    #   BM logo  — anchored: ~1.67" wide x 1.19" tall
    #              H offset from column: ~5.28" (4,833,620 EMU)
    #              V offset from paragraph: -0.14" (-123,825 EMU)
    # =========================================================
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if nc_logo_path and os.path.exists(nc_logo_path):
        logo_para.add_run().add_picture(nc_logo_path, height=Inches(0.9))
    else:
        add_run(logo_para, "NIPPLE CRIME", bold=True, size=20)

    if bm_logo_path and os.path.exists(bm_logo_path):
        add_floating_image(
            logo_para,
            bm_logo_path,
            width_emu=1529080,
            height_emu=1091565,
            pos_h_emu=4833620,
            pos_v_emu=-123825,
        )

    doc.add_paragraph()  # spacer

    # =========================================================
    # DIVIDER LINE
    # =========================================================
    div = doc.add_paragraph()
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # =========================================================
    # TITLE BLOCK
    # =========================================================
    label_para = doc.add_paragraph()
    label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(label_para, "STANDARD OPERATING PROCEDURE", bold=True, size=11)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title_para, sop_title, bold=True, size=15)

    doc.add_paragraph()

    # =========================================================
    # METADATA TABLE
    # =========================================================
    meta_table = doc.add_table(rows=2, cols=4)
    meta_table.style = 'Table Grid'

    meta_rows = [
        [("SOP Number", sop_number), ("Department", department)],
        [("Version", version),       ("Effective Date", effective_date)],
    ]

    for r, row_data in enumerate(meta_rows):
        row = meta_table.rows[r]
        for c, (lbl, val) in enumerate(row_data):
            label_cell = row.cells[c * 2]
            value_cell = row.cells[c * 2 + 1]
            set_cell_bg(label_cell, "D9D9D9")
            lp = label_cell.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(lp, lbl, bold=True, size=9)
            vp = value_cell.paragraphs[0]
            add_run(vp, val, size=9)

    lu_row = meta_table.add_row()
    merged = lu_row.cells[0].merge(lu_row.cells[1]).merge(lu_row.cells[2]).merge(lu_row.cells[3])
    set_cell_bg(merged, "D9D9D9")
    lup = merged.paragraphs[0]
    lup.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(lup, "Last Updated: %s" % last_updated, bold=True, size=9)

    doc.add_paragraph()

    # =========================================================
    # BODY SECTIONS
    # =========================================================
    for (level, heading, lines) in sections:
        h = doc.add_heading(heading, level=level)
        h.runs[0].font.size = Pt({1: 13, 2: 11, 3: 10}.get(level, 10))
        h.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        for line in lines:
            if line.startswith("- "):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif len(line) > 2 and line[0].isdigit() and line[1] == '.':
                p = doc.add_paragraph(line[line.index(".")+1:].strip(), style='List Number')
            elif len(line) > 3 and line[:2].isdigit() and line[2] == '.':
                p = doc.add_paragraph(line[line.index(".")+1:].strip(), style='List Number')
            elif line == "":
                p = doc.add_paragraph()
                continue
            else:
                p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(10)

    # =========================================================
    # FOOTER
    # =========================================================
    footer = page_section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp,
        "Nipple Crime Theme Camp  |  %s %s  |  Ver %s  |  Confidential — Internal Use Only"
        % (sop_number, sop_title, version),
        size=8, color=(128, 128, 128))

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print("Saved: %s" % output_path)


# =========================================================
# SOP Tr3 — Statement of Intent (BMorg)
# NOTE: Tr3 master template was manually finalized. Do not regenerate.
# =========================================================

# =========================================================
# SOP Tr4 — Mutant Vehicle Statement of Intent (BMorg)
# =========================================================

sections_tr4 = [
    (1, "1. Purpose", [
        "This SOP documents the annual process for completing and submitting Nipple Crime's "
        "Mutant Vehicle Statement of Intent (MVSOI) to the Burning Man Organization (BMorg). "
        "The MVSOI is required for all mutant vehicles seeking to operate on the playa and to "
        "be considered for the Stewards Ticket Sale allocation. Primary vehicle: SirKiss.",
    ]),
    (1, "2. Timeline", [
        "- MVSOI portal opens: December – January (check burningman.org each year)",
        "- Submission deadline: February – March (verify annually)",
        "- Stewards Sale ticket allocations announced: second week of February",
        "- On-playa DMV licensing: upon arrival at Burning Man each year",
        "- Set calendar reminders at each milestone",
    ]),
    (1, "3. Submission Steps", [
        "1. Go to burningman.org > Participate > Mutant Vehicles > Statement of Intent",
        "2. Log in using the Nipple Crime account (credentials held by President and Treasurer)",
        "3. Select 'Previously Applied — We've applied to bring this vehicle before'",
        "4. Complete all required fields — see Section 4 for 2026 reference answers",
        "5. Review entries before submitting",
        "6. Submit form",
        "7. Save / screenshot the BMorg confirmation email for records (see Section 5)",
    ]),
    (1, "4. 2026 Submission Reference", [
        "The following answers were submitted for the 2026 Burn cycle. Update each year as needed.",
        "",
        "CONTACT",
        "- First Name: Reece",
        "- Last Name: Dassinger",
        "- Email: Leadership@nipplecrime.org",
        "",
        "VEHICLE",
        "- Primary Mutant Vehicle Name: SirKiss",
        "- Most Recent Placed Camp Name: Nipple Crime",
        "- Vehicle Status: Previously Applied — We've applied to bring this vehicle before",
        "- Most Recent Year Licensed at On-Playa DMV: 2025",
        "",
        "INTERACTIVITY",
        "- Participatory Aspects: Offering Rides, Music/Sound System, DJ Platform",
        "- Description: We have a sound tech, DJ lineups, crowd management, walkers, and we allow "
        "anyone on the art car at all times. We are handicap accessible. We have been in the zip line.",
        "",
        "STATUS FOR 2026",
        "- Active / Requesting Access: We plan to bring our Mutant Vehicle in 2026 and would like "
        "to be considered for the Stewards Ticket Sale this year.",
        "",
        "MULTIPLE VEHICLES",
        "- Registering more than one MV: No",
        "",
        "TICKET INFORMATION",
        "- Total crew and passenger capacity of all MVs: 60",
        "- Tickets requested for MV Crew and Support Team: 18",
        "  (Note: tickets sold in pairs; MV crew only — theme camp support tickets requested separately through Placement)",
        "",
        "CAMPING PLANS",
        "- Part of Another Placed Camp: Team members camping with Nipple Crime (placed theme camp "
        "submitting its own placement request)",
    ]),
    (1, "5. Record Keeping", [
        "- Save BMorg confirmation email as PDF",
        "- File location: [Shared Drive > Treasurer > BMorg > MVSOI > YYYY]",
        "- Log: submission date, submitted by, BMorg correspondence, ticket allocation result",
        "- Retain records for minimum 5 years per nonprofit compliance requirements",
    ]),
    (1, "6. Contacts", [
        "- BMorg Mutant Vehicle Team: burningman.org (contact via MV portal)",
        "- President / Primary Submitter: Reece Dassinger — reece@nipplecrime.org",
        "- Art Car Supervisor (SirKiss): Anthony Tolosano",
        "- Treasurer: Isabel Hoy — izhoy@yahoo.com",
        "- VP / SOP Owner: Chris Reddin — creddin1@hotmail.com",
    ]),
    (1, "7. Revision History", [
        "- v1.0 | 2026-03-03 | Initial draft | Chris Reddin",
    ]),
]

create_sop(
    output_path="Standard Operating Procedures/Tr4 Mutant Vehicle Statement of Intent.docx",
    sop_number="Tr4",
    sop_title="Mutant Vehicle Statement of Intent (BMorg)",
    department="Treasurer",
    version="1.0",
    effective_date="2026-03-03",
    last_updated="2026-03-03",
    sections=sections_tr4,
)

# =========================================================
# SOP Com1 — Slack
# =========================================================

sections_com1 = [
    (1, "1. Purpose", [
        "This SOP covers how Nipple Crime uses Slack as its central internal communications "
        "platform. It defines the channel structure, posting rules, and the meeting recap "
        "process managed by the Communications Officer.",
    ]),
    (1, "2. Channel Structure", [
        "- #announcements — Camp-wide announcements; replies in thread only",
        "- #leadership-board — Board-level discussion (board members only)",
        "- #leadership-general — All leadership roles",
        "- #committee-food — Kitchen & Bar committee",
        "- #committee-power — Power Grid team",
        "- #committee-lnt — Leave No Trace team",
        "- #help-wanted — Open asks and volunteer recruitment",
        "Additional committee channels created as needed using the #committee- prefix.",
    ]),
    (1, "3. Posting Rules", [
        "- Announcements: one top-level post, all replies in thread",
        "- Decisions: summarize in one 'Decision:' message and pin it to the channel",
        "- Pin the following in relevant channels at all times:",
        "  - Interest / labor form link",
        "  - Dues payment link (when available)",
        "  - Org chart",
        "  - Meeting recap index",
    ]),
    (1, "4. Meeting Recap SOP", [
        "Owner: Communications Officer",
        "Deadline: Post within 24 hours of any board or leadership meeting.",
        "Post to: #leadership-board in Slack (and optionally email the board).",
        "",
        "TEMPLATE (copy/paste each recap):",
        "",
        "Meeting: Board Meeting -- YYYY-MM-DD",
        "Attendees: [list names]",
        "",
        "Key Decisions:",
        "- Decision 1 (who decided, any constraints)",
        "- Decision 2",
        "",
        "Action Items (Owner | Action | Due Date | Dependencies):",
        "- [Owner] -- [action] -- [due date] -- [dependencies]",
        "",
        "Risks / Watch-outs:",
        "- [Risk] -- [mitigation owner] -- [next check-in date]",
        "",
        "Next Meeting:",
        "- Date/time + what must be ready by then",
    ]),
    (1, "5. Contacts", [
        "- Communications Officer (Slack owner): Daemon Wyner",
        "- President (approves major messaging): Reece Dassinger",
        "- VP / SOP Owner: Chris Reddin",
    ]),
    (1, "6. Revision History", [
        "- v1.0 | 2026-03-03 | Initial draft | Chris Reddin",
    ]),
]

create_sop(
    output_path="Standard Operating Procedures/Com1 Slack.docx",
    sop_number="Com1",
    sop_title="Slack",
    department="Communications",
    version="1.0",
    effective_date="2026-03-03",
    last_updated="2026-03-03",
    sections=sections_com1,
)

# =========================================================
# SOP Com2 — Sakari
# =========================================================

sections_com2 = [
    (1, "1. Purpose", [
        "This SOP covers how Nipple Crime uses Sakari for SMS outreach to camp members. "
        "Sakari is integrated with HubSpot, allowing bulk and 1:1 texts to be sent "
        "directly to HubSpot contact lists.",
    ]),
    (1, "2. Access & Integration Setup", [
        "- Confirm Sakari is connected to HubSpot (Settings > Integrations in Sakari)",
        "- When connected, HubSpot contact lists sync and a Sakari SMS card appears on each contact record",
        "- Credentials held by: Communications Officer and President",
    ]),
    (1, "3. Compliance (Opt-Out)", [
        "- Opt-out is carrier keyword-driven: members text STOP to opt out",
        "- Always include 'Reply STOP to opt out' on first outreach of each year",
        "- Do not override carrier opt-out behavior",
        "- Pricing is driven by SMS segments (longer messages or special characters cost more -- keep messages short)",
    ]),
    (1, "4. Send a Bulk SMS Campaign", [
        "Use this to text the whole membership (e.g., annual interest check-in).",
        "",
        "1. In Sakari, go to Campaigns > Create Campaign",
        "2. Complete the campaign stages: Details, Contacts, Conditions, Messaging, Schedule",
        "3. In Contacts, select the synced HubSpot list (e.g., '2026 - Interest: Yes/Maybe/Unknown')",
        "4. Write message -- keep it short, include the form link",
        "5. Preview estimated cost: three-dot menu > Preview before sending",
        "6. Schedule at a reasonable hour; avoid repeat sends to non-responders too quickly",
        "7. Send",
        "",
        "RECOMMENDED TEMPLATE (first outreach of year):",
        "NC 2026 check-in: are you coming + can you help? Fill this out: [FORM LINK]",
        "Reply STOP to opt out.",
    ]),
    (1, "5. Send a 1:1 SMS (from HubSpot)", [
        "Use this for personal follow-up with potential leaders or unclear respondents.",
        "",
        "1. In HubSpot, open the person's Contact record",
        "2. Find the Sakari SMS card/module",
        "3. Click Send SMS, write your message, send",
        "4. Replies are tracked on the contact record",
        "",
        "RECOMMENDED TEMPLATE (role recruitment):",
        "Hey -- you marked that you're open to leadership. We still need a [Role] Lead.",
        "Are you open to it (or co-leading)? I can send a simple checklist.",
        "",
        "Troubleshooting: If the SMS card is missing, confirm the Sakari-HubSpot integration",
        "is active and the card is enabled in HubSpot's integration settings.",
    ]),
    (1, "6. Contacts", [
        "- Communications Officer (Sakari owner): Daemon Wyner",
        "- President: Reece Dassinger",
        "- VP / SOP Owner: Chris Reddin",
    ]),
    (1, "7. Revision History", [
        "- v1.0 | 2026-03-03 | Initial draft | Chris Reddin",
    ]),
]

create_sop(
    output_path="Standard Operating Procedures/Com2 Sakari.docx",
    sop_number="Com2",
    sop_title="Sakari (SMS)",
    department="Communications",
    version="1.0",
    effective_date="2026-03-03",
    last_updated="2026-03-03",
    sections=sections_com2,
)

# =========================================================
# SOP Com3 — HubSpot
# =========================================================

sections_com3 = [
    (1, "1. Purpose", [
        "This SOP covers how Nipple Crime uses HubSpot for contact management, forms, "
        "email marketing (Nipple News), and website upkeep. HubSpot is the central CRM "
        "and outbound communications platform managed by the Communications Officer.",
    ]),
    (1, "2. Access & Permissions", [
        "Before doing anything else, verify you have the correct permissions.",
        "- Minimum required: Super Admin (or equivalent) access",
        "- Required access areas:",
        "  - CRM (Contacts)",
        "  - Marketing > Forms",
        "  - Marketing > Email",
        "  - Content / Website Pages",
        "  - Commerce (optional; usually Treasurer-owned)",
        "- To check: Settings > Users & Teams in HubSpot",
        "- If you cannot create/edit properties, lists, or forms -- fix permissions before troubleshooting anything else",
        "- Note: Property-level access can be restricted separately from general contact access",
    ]),
    (1, "3. Data Model", [
        "- Every person is a Contact record",
        "- Contact properties store structured answers (e.g., 'Interested in camp 2026?', 'Volunteer hours', 'Leadership role')",
        "- Forms collect and update properties at scale",
        "- Lists/Segments (active or static) target the right people for outreach",
    ]),
    (1, "4. Naming Conventions", [
        "Use consistent naming so properties and lists stay organized across years.",
        "",
        "PROPERTIES (prefix with year):",
        "- 2026 - Interested in Camp",
        "- 2026 - Volunteer Time Available",
        "- 2026 - Leadership Interest",
        "- 2026 - Leadership Role (Assigned)",
        "",
        "LISTS/SEGMENTS:",
        "- 2026 - Interest: Yes/Maybe/Unknown",
        "- 2026 - Interest: No",
        "- 2026 - Potential Leaders (Interested, No Role Assigned)",
        "- 2026 - Needs Follow-up (No response yet)",
        "",
        "List type rule:",
        "- Active lists: anything ongoing (auto-updates as responses come in)",
        "- Static lists: one-time snapshots (everyone invited to X at time of send)",
    ]),
    (1, "5. Create a Contact Property", [
        "Use this to add a structured field (e.g., a dropdown) for survey answers.",
        "",
        "1. In HubSpot, click Settings (gear icon)",
        "2. Go to Data Management > Properties",
        "3. Choose object type: Contact",
        "4. Click Create Property",
        "5. Set:",
        "   - Property label: e.g., 2026 - Interested in Camp",
        "   - Field type: Dropdown select",
        "   - Options: Yes, Maybe, No",
        "6. Save",
        "",
        "The property is now available in forms, contact views, and list filters.",
        "Troubleshooting: If 'Create property' is missing or throws errors, check Super Admin permissions.",
    ]),
    (1, "6. Build the Annual Interest / Labor / Leadership Form", [
        "Clone last year's form rather than building from scratch.",
        "",
        "1. Go to Marketing > Forms",
        "2. Find last year's 'Existing Member Interest' form",
        "3. Clone it",
        "4. Update text to 2026; keep it short",
        "5. Required fields:",
        "   - Email (required)",
        "   - 2026 - Interested in Camp (dropdown: Yes / Maybe / No)",
        "   - 2026 - Volunteer Time Available (dropdown)",
        "   - 2026 - Leadership Interest (dropdown)",
        "6. Recommended dropdown options for Volunteer Time Available:",
        "   - 1-2 hours/week",
        "   - 3-4 hours/week",
        "   - 4+ hours/week",
        "   - I have time mainly close to the burn",
        "   - No time / can't help this year",
        "7. Recommended dropdown options for Leadership Interest:",
        "   - Yes -- I'm interested in a leadership role",
        "   - Maybe -- talk to me",
        "   - No -- not this year",
        "8. Publish the form",
        "9. Get the share link: Marketing > Forms > hover form > Actions > Share > Copy link",
        "10. Quality check: submit it yourself once and confirm the properties updated on your contact record",
    ]),
    (1, "7. Create a Segment / List for Targeting", [
        "1. Go to CRM > Lists (Segments)",
        "2. Create a Contact-based list",
        "3. Choose list type (Active or Static -- see Section 4)",
        "4. Build filters",
        "",
        "EXAMPLE A: '2026 - Interest: Yes/Maybe/Unknown' (Active list)",
        "Filter: 2026 - Interested in Camp is any of: Yes, Maybe",
        "(and/or is unknown/empty for non-responders)",
        "",
        "EXAMPLE B: 'Potential Leaders (Willing + No Role Assigned)'",
        "Prereq: create '2026 - Leadership Role (Assigned)' property first",
        "Filter 1: 2026 - Leadership Interest is any of: Yes, Maybe",
        "Filter 2: AND 2026 - Leadership Role (Assigned) is unknown/empty",
        "Save as: '2026 - Potential Leaders (No Role Assigned)'",
    ]),
    (1, "8. Send Nipple News Email", [
        "Clone last year's closest matching email rather than building from scratch.",
        "",
        "1. Go to Marketing > Email",
        "2. Find last year's matching email (e.g., 'Are you coming?')",
        "3. Clone",
        "4. Edit: subject line (clear + action-oriented), body (short sections, bold deadlines), insert form link",
        "5. Preview as a real contact; send a test email to yourself",
        "6. Set Send to list: '2026 - Interest: Yes/Maybe/Unknown'",
        "7. Enable web version ('view in browser') so you can share the link on Facebook or via SMS",
        "8. Send",
        "",
        "WEB VERSION SHARE TEMPLATE (SMS/Facebook):",
        "Nipple News #1 is out -- full details here: [WEB VERSION LINK]",
    ]),
    (1, "9. Get the Web Version Link", [
        "1. Open the email in HubSpot",
        "2. Go to email details / web version area",
        "3. Copy the web version URL",
        "Use this link for: Facebook group posts, SMS 'read more' links, members who don't read email.",
    ]),
    (1, "10. Website Upkeep (HubSpot CMS)", [
        "Keep Sign Up, Donate, and other key pages current each year.",
        "",
        "1. Go to Content > Website Pages",
        "2. Click the page to edit",
        "3. Update text, dates, and buttons",
        "4. Publish (or schedule publish)",
        "",
        "If temporarily disabling a button (e.g., donations not open yet):",
        "- Remove the button URL, OR",
        "- Change the button URL to the homepage to avoid a dead link",
    ]),
    (1, "11. Payment Links (Know Where They Live)", [
        "Payment links are owned by the Treasurer but Comms touches them when updating pages or emails.",
        "Location: HubSpot > Commerce > Payment Links",
        "",
        "Comms responsibilities:",
        "- Update broken or mislabeled buttons on web pages (notify Treasurer when doing so)",
        "- Do NOT change pricing or payment plan logic without explicit Treasurer approval",
    ]),
    (1, "12. Handoff Checklist (for Next Comms Officer)", [
        "Keep this updated throughout the year so handoff is smooth.",
        "",
        "- List of active properties created this year (names + purpose)",
        "- List/segment names used",
        "- Which form is the canonical interest form for this year",
        "- Message templates:",
        "  - First interest SMS",
        "  - Leadership recruitment SMS",
        "  - Nipple News email skeleton",
        "- Where web version links live (and how to generate them)",
        "- Website pages most often updated (Sign Up, Donate, etc.)",
        "- Who to contact for permissions: President / Super Admin",
    ]),
    (1, "13. Contacts", [
        "- Communications Officer (HubSpot owner): Daemon Wyner",
        "- President (approves major messaging + payment decisions): Reece Dassinger",
        "- Treasurer (payment links, dues language): Isabel Hoy",
        "- VP / SOP Owner: Chris Reddin",
    ]),
    (1, "14. Revision History", [
        "- v1.0 | 2026-03-03 | Initial draft | Chris Reddin",
    ]),
]

# =========================================================
# SOP Tr2 — IRS 509(a)(2) Administration
# =========================================================

sections_tr2 = [
    (1, "1. Purpose", [
        "This SOP documents the annual compliance requirements for maintaining Nipple Crime's "
        "federal tax-exempt status under IRC Section 501(c)(3) with public charity classification "
        "under Section 509(a)(2). Compliance is owned by the Treasurer with oversight from the President.",
    ]),
    (1, "2. IRS Determination Letter — Key Details", [
        "The original IRS determination letter (dated 05/06/2024) is the foundational legal document "
        "for the organization's nonprofit status. Key details from that letter:",
        "",
        "- Legal Name: Camp NC Inc",
        "- Address on Record: 527 S Arlington Ave, Reno, NV 89509",
        "- Employer ID Number (EIN): 99-1556119",
        "- Public Charity Status: 509(a)(2)",
        "- Tax-Exempt Under: IRC Section 501(c)(3)",
        "- Effective Date of Exemption: February 23, 2024",
        "- Accounting Period Ending: January 31 (fiscal year end)",
        "- Form 990 / 990-EZ / 990-N Required: Yes (annually)",
        "- Contribution Deductibility: Yes -- donors may deduct contributions under IRC Section 170",
        "- Addendum Applies: No",
        "- IRS Document Locator Number (DLN): 26053516003724",
        "- IRS Contact: Customer Service, ID 31954, (877) 829-5500",
        "",
        "The original letter is filed at: [Shared Drive > Treasurer > IRS > Determination Letter 2024]",
        "A copy should also be kept with the Secretary/President.",
    ]),
    (1, "3. Annual Filing Requirements", [
        "Fiscal year end: January 31",
        "Form 990 series due date: June 15 (15th day of the 5th month after fiscal year end)",
        "Extension available: 6-month automatic extension to December 15 (file Form 8868 before June 15)",
        "",
        "Which Form 990 to file (based on gross receipts):",
        "- Form 990-N (e-Postcard): gross receipts normally $50,000 or less",
        "- Form 990-EZ: gross receipts $50,001 - $200,000 AND total assets under $500,000",
        "- Form 990 (full): gross receipts over $200,000 OR total assets $500,000 or more",
        "",
        "Treasurer determines which form applies each year based on financials.",
        "Filing is done at IRS.gov (tax-exempt organizations section) or through a CPA.",
    ]),
    (1, "4. Annual Compliance Checklist", [
        "Complete the following each fiscal year (year end: January 31):",
        "",
        "FEBRUARY (immediately after fiscal year end):",
        "- Close the books for the fiscal year ended January 31",
        "- Compile income and expense summary by category",
        "- Reconcile all bank accounts",
        "- Collect receipts and documentation for all expenses",
        "",
        "MARCH - MAY:",
        "- Determine which Form 990 applies (see Section 3)",
        "- Prepare or engage CPA to prepare the return",
        "- Board reviews draft Form 990 before filing",
        "- President signs the return",
        "",
        "BY JUNE 15:",
        "- File Form 990 / 990-EZ / 990-N with IRS",
        "- OR file Form 8868 for 6-month extension (if needed)",
        "- Save filed return and IRS acknowledgment to: [Shared Drive > Treasurer > IRS > 990 > YYYY]",
        "",
        "ONGOING:",
        "- Notify IRS of any change in address or responsible party (Form 8822-B within 60 days)",
        "- Do not allow status to lapse -- failure to file 3 consecutive years results in automatic revocation",
        "- Keep determination letter accessible at all times (required to show donor deductibility)",
    ]),
    (1, "5. What 509(a)(2) Means for Camp Operations", [
        "- Nipple Crime is classified as a publicly supported charity (not a private foundation)",
        "- Donors may deduct contributions on their federal tax returns",
        "- Camp must maintain broad public support -- dues, ticket sales, and participation fees "
        "generally qualify as public support under 509(a)(2)",
        "- Excessive reliance on a single donor or source could jeopardize status -- flag to Treasurer if any "
        "single donor contributes more than 33% of total annual revenue",
        "- Investment income is permissible but should remain incidental",
    ]),
    (1, "6. Record Keeping", [
        "Retain all of the following for a minimum of 7 years:",
        "- IRS determination letter",
        "- All filed Form 990 / 990-EZ / 990-N returns",
        "- Form 8868 extension requests (if any)",
        "- Annual financial statements and bank reconciliations",
        "- Board meeting minutes approving annual budget and major expenditures",
        "- Any IRS correspondence",
        "",
        "File location: [Shared Drive > Treasurer > IRS > ]",
    ]),
    (1, "7. Contacts", [
        "- IRS Tax-Exempt Organizations: (877) 829-5500 (Customer Service ID: 31954)",
        "- IRS Website: irs.gov/charities-non-profits",
        "- Treasurer (filing owner): Isabel Hoy -- izhoy@yahoo.com",
        "- President (signs returns, oversight): Reece Dassinger -- reece@nipplecrime.org",
        "- VP / SOP Owner: Chris Reddin -- creddin1@hotmail.com",
    ]),
    (1, "8. Revision History", [
        "- v1.0 | 2026-03-04 | Initial draft | Chris Reddin",
    ]),
]

create_sop(
    output_path="Standard Operating Procedures/Tr2 IRS 509a2 Administration.docx",
    sop_number="Tr2",
    sop_title="IRS 509(a)(2) Administration",
    department="Treasurer",
    version="1.0",
    effective_date="2026-03-04",
    last_updated="2026-03-04",
    sections=sections_tr2,
)

create_sop(
    output_path="Standard Operating Procedures/Com3 HubSpot.docx",
    sop_number="Com3",
    sop_title="HubSpot",
    department="Communications",
    version="1.0",
    effective_date="2026-03-03",
    last_updated="2026-03-03",
    sections=sections_com3,
)

# =========================================================
# SOP HR3 — Stewards Sale (BMorg Ticket Allocation)
# =========================================================

sections_hr3 = [
    (1, "1. Purpose", [
        "This SOP documents the annual process for managing Nipple Crime's Burning Man "
        "Stewards Sale ticket allocation. The Stewards Sale is a directed ticket sale run by "
        "the Burning Man Organization (BMorg) for placed theme camps and registered mutant "
        "vehicles. Camps receive a limited block of one-time-use ticket codes which are "
        "distributed to returning members in good standing. This SOP defines internal "
        "eligibility, distribution mechanics, and recordkeeping. See Tr4 for the related "
        "Mutant Vehicle Statement of Intent process.",
    ]),
    (1, "2. Timeline", [
        "- Placement / MVSOI submitted: February (see Tr3, Tr4)",
        "- Stewards Sale allocations announced by BMorg: second week of February",
        "- Camp receives ticket codes from BMorg: late February (date set by BMorg each year)",
        "- Internal eligibility list finalized by HR: within 5 business days of allocation announcement",
        "- Codes distributed to members: within 3 business days of receipt from BMorg",
        "- Member redemption window (BMorg portal): typically 7-10 days; exact window set by BMorg",
        "- Unredeemed code reclamation: 48 hours before redemption window closes",
        "- Set calendar reminders at each milestone — the window is tight and codes do not roll over",
    ]),
    (1, "3. Roles & Responsibilities", [
        "The Stewards Sale is owned end-to-end by Human Resources. No camp funds move "
        "through this process — members purchase tickets directly from BMorg using "
        "personally-paid funds — so the Treasurer has no operational role here.",
        "",
        "HUMAN RESOURCES (full ownership):",
        "- Maintain the Returning Member roster and good-standing status year-round",
        "- Serve as Nipple Crime's point of contact with BMorg for the Stewards Sale",
        "- Receive the ticket code block from BMorg and confirm count matches the allocation",
        "- Produce and finalize the internal eligibility list each Stewards Sale cycle",
        "- Distribute codes to eligible members (see Section 5)",
        "- Communicate eligibility, deadlines, and redemption instructions",
        "- Track redemption status and reclaim/reassign unredeemed codes (see Section 6)",
        "- File records per Section 7",
        "",
        "PRESIDENT (oversight):",
        "- Final approval on the eligibility list and any exceptions",
        "- Resolve disputes; sign off on reclamations and reassignments",
        "",
        "MEMBERS (individual responsibility):",
        "- Pay BMorg directly for their ticket using the assigned code",
        "- Confirm redemption to HR within 24 hours of purchase",
        "- Carry full personal responsibility for their own ticket funds; the camp does "
        "not collect, hold, or refund ticket payments",
    ]),
    (1, "4. Member Eligibility — Returning Members in Good Standing", [
        "Stewards Sale codes are distributed only to Returning Members in Good Standing. A "
        "member meets this standard if ALL of the following are true at the time the "
        "eligibility list is finalized:",
        "",
        "- Camped with Nipple Crime in at least one of the two prior Burns",
        "- Dues for the current Burn year are paid (or on a Treasurer-approved payment plan)",
        "- Met or committed to the camp's minimum volunteer hours for the current year",
        "- No outstanding incident, code-of-conduct, or LNT issues from prior Burns",
        "- Confirmed 'Yes' on the current year's interest form (see Com3)",
        "",
        "PRIORITY ORDER (when codes are fewer than eligible members):",
        "1. Active Board members and Leadership role holders",
        "2. Confirmed critical-crew leads (Build, Kitchen, Power, LNT, Medical, Safety)",
        "3. Returning members with 2+ prior Burns and dues paid in full",
        "4. Returning members with 1 prior Burn and dues paid in full",
        "5. All other Returning Members in Good Standing",
        "",
        "Exceptions to eligibility (e.g., first-year leadership, hardship cases) require "
        "President approval and must be logged in the eligibility record.",
    ]),
    (1, "5. Code Distribution Process", [
        "Members purchase tickets directly from BMorg using one-time-use codes issued to "
        "the camp. Nipple Crime does NOT collect ticket payments or purchase tickets on behalf "
        "of members.",
        "",
        "1. HR receives the code block from BMorg via email or BMorg portal",
        "2. HR confirms the code count matches the allocation",
        "3. HR confirms the finalized eligibility list with the President",
        "4. HR assigns one code per eligible member, in priority order, until codes are exhausted",
        "5. HR emails each member individually with:",
        "   - Their assigned code (do NOT post codes in Slack or any shared channel)",
        "   - The redemption deadline (date and time, with timezone)",
        "   - A direct link to the BMorg redemption portal",
        "   - A reminder that codes are one-time use and tied to the member's BMorg account",
        "   - Instructions to reply confirming redemption within 24 hours of purchase",
        "6. HR logs each code-to-member assignment in the Stewards Sale tracker",
        "7. Members redeem in the BMorg portal during the redemption window",
        "",
        "SECURITY: Codes have monetary value. Treat them like cash.",
        "- Never post codes in any Slack channel, public document, or group email",
        "- Send codes only via direct 1:1 email to the assigned member",
        "- If a code is leaked or compromised, HR notifies BMorg and the President immediately",
    ]),
    (1, "6. Unredeemed Code Reclamation", [
        "Unused codes are forfeited at the end of the BMorg redemption window. To avoid waste:",
        "",
        "- 72 hours before close: HR sends a reminder to any member who has not confirmed redemption",
        "- 48 hours before close: HR reclaims any code where the member has not confirmed or responded",
        "- Reclaimed codes are reassigned to the next eligible member on the priority list (Section 4)",
        "- Reassignments follow the same 1:1 email process from Section 5",
        "- All reclaim/reassign actions require President sign-off and are logged in the tracker",
    ]),
    (1, "7. Record Keeping", [
        "Maintain the following each Stewards Sale cycle. Retain records for a minimum of 5 years.",
        "",
        "- Eligibility list (final), with priority tier and exceptions noted",
        "- Code-to-member assignment log (member name, email, code last 4, assignment date, redemption status)",
        "- Member redemption confirmations (forwarded emails or screenshots)",
        "- BMorg correspondence (allocation announcement, code block delivery, post-sale report)",
        "- Reclamation/reassignment log with President approvals",
        "",
        "File location: [Shared Drive > HR > Stewards Sale > YYYY]",
        "Do NOT retain the full ticket codes themselves after the redemption window closes — store only last 4 digits for reconciliation.",
        "Note: No financial records are produced by this process. Member ticket payments are personal "
        "transactions with BMorg and are not part of camp bookkeeping.",
    ]),
    (1, "8. Contacts", [
        "- HR Officer (process owner, BMorg POC, code distribution): [TBD] — [email]",
        "- President (approvals + exceptions): Reece Dassinger — reece@nipplecrime.org",
        "- VP / SOP Owner: Chris Reddin — creddin1@hotmail.com",
        "- BMorg Stewards Sale: tickets.burningman.org / tickets@burningman.org",
    ]),
    (1, "9. Revision History", [
        "- v1.0 | 2026-05-13 | Initial draft | Chris Reddin",
    ]),
]

create_sop(
    output_path="Standard Operating Procedures/HR3 Stewards Sale.docx",
    sop_number="HR3",
    sop_title="Stewards Sale (BMorg Ticket Allocation)",
    department="Human Resources",
    version="1.0",
    effective_date="2026-05-13",
    last_updated="2026-05-13",
    sections=sections_hr3,
)

# =========================================================
# SOP Tr5 — Hell Station Petrol Application (BMorg)
# =========================================================

sections_tr5 = [
    (1, "1. Purpose", [
        "This SOP documents the annual process for registering Nipple Crime with the "
        "Burning Man Organization (BMorg) PETROL Department for the Black Rock City (BRC) "
        "Fuel Program. Successful registration enables on-playa diesel delivery to the "
        "camp generator(s) and propane refueling for kitchen, bar, and any registered "
        "flame effects. The program is intended to reduce unsafe fuel transport into BRC "
        "and consolidate fuel storage. Pre-event registration is owned by the Treasurer; "
        "on-playa account activation and fuel operations are owned by the Infrastructure "
        "Director. Hell Station is located at the green dot at roughly 9:30 & halfway to "
        "fenceline.",
    ]),
    (1, "2. Timeline", [
        "- Fuel registration window opens (BMorg): February 15",
        "- Camp delivery application deadline (firm — 2026): March 31, 2026",
        "- Fuel registration window closes (BMorg): June 15",
        "- Account approval + payment activation email from BMorg/Gravity: after July 1",
        "- Payment account setup (Gravity, save credit card on auto-pay): early August",
        "- Hell Station Office opens (2025 reference): August 18 — verify 2026 dates each year",
        "- Hell Station fuel lanes open (2025 reference): August 19",
        "- Hell Station propane opens (2025 reference): August 20",
        "- Hell Station closes (2025 reference): August 31",
        "- On-playa account activation: as soon as Infrastructure Director arrives with gov't ID",
        "- First fuel delivery: minimum 24 hours after account activation + asset registration",
        "- Drum pickup window: starts Sunday Temple Burn day; final drop-off at Hell Station Monday 4:00 pm",
        "- There are no on-playa registrations except for DMV-registered accessibility vehicles — miss the window and we get no fuel",
        "- Set calendar reminders at each milestone",
    ]),
    (1, "3. Roles & Responsibilities", [
        "Ownership splits at the playa gate: Treasurer runs the pre-event registration and "
        "payment side; Infrastructure Director runs everything on-playa.",
        "",
        "TREASURER (pre-event ownership):",
        "- Submit the BRC Fuel Program application by the March 31 camp delivery deadline",
        "- Coordinate with Placement (via Tr3 lead) so fuel delivery location aligns with placed camp layout",
        "- Set up the PETROL Payment Account in Gravity once approval email arrives (after July 1)",
        "- Save the camp credit card on auto-pay; keep card active through post-event reconciliation",
        "- Designate account users (Project Lead, Fuel Contact, Payment Contact) — see Section 5",
        "- Hand off account credentials and Customer Portal access to the Infrastructure Director before the burn",
        "",
        "INFRASTRUCTURE DIRECTOR (on-playa ownership):",
        "- Activate the account at Hell Station with government-issued ID on arrival",
        "- Host the Field Registration Agents when they visit camp to tag approved assets",
        "- Ensure camp fuel storage and generator placement meet Section 7 requirements before agents arrive",
        "- Manage day-to-day deliveries, drum receipt, propane pickup, and the spill kit",
        "- Coordinate drum pickup at end of event (or drop at Hell Station by Mon 4:00 pm)",
        "- Report any spill, leak, or compromised containment to PETROL, Rangers, or ESD immediately",
        "",
        "PRESIDENT (oversight):",
        "- Approve fuel budget and any in-season changes",
        "- Resolve disputes with BMorg/PETROL if escalated",
        "",
        "PAYMENT CONTACT (must be on playa for activation):",
        "- The cardholder must be present at Hell Station to activate the account, OR",
        "- Treasurer must notify PETROL in advance if the cardholder will not be on playa",
        "- Default: Treasurer is Payment Contact; if Treasurer is not on playa, document arrangement with PETROL during registration",
    ]),
    (1, "4. Application Submission (Pre-Event)", [
        "Apply through the participant's Burner Profile each year. The application captures "
        "camp layout, equipment, and estimated fuel needs; it triggers a safety review and "
        "may require layout revisions before approval.",
        "",
        "1. Log in at profiles.burningman.org",
        "2. Go to Participation > BRC Fuel Program application",
        "3. Submit BEFORE March 31 to qualify for in-camp delivery (Placement coordination)",
        "4. Complete all sections — see Section 5 for the required data",
        "5. Upload required diagrams — see Section 5",
        "6. Submit and save the BMorg confirmation email/PDF to records (Section 11)",
        "7. Watch for follow-up from PETROL — they may require camp layout revisions for safety/delivery access; respond promptly",
        "8. Confirmation of approval arrives via email; payment activation link follows after July 1",
        "",
        "If we miss March 31, the application stays open through June 15 but BMorg cannot guarantee delivery — we'd have to fall back to Hell Station pickup only.",
    ]),
    (1, "5. Application Data — Nipple Crime Reference", [
        "Use these answers as the 2026 baseline. Update each year as needed.",
        "",
        "ACCOUNT USERS (designate at least one per role; one person can fill multiple):",
        "- Project Lead: Infrastructure Director — [Name TBD] — [email]",
        "- Fuel Contact: Infrastructure Director — [Name TBD] — [email]",
        "- Payment Contact: Treasurer — Isabel Hoy — izhoy@yahoo.com",
        "",
        "CAMP INFO:",
        "- Camp Name: Nipple Crime",
        "- Camp Type: Placed Theme Camp (confirmed via Placement — see Tr3)",
        "- Camp size and frontage: pulled from current year's Placement submission",
        "",
        "FUEL NEEDS (confirm with Infrastructure team before submitting each year):",
        "- Diesel (dyed): for camp generator(s) — request delivery to generator location, or 55-gal drum delivery (see Section 8)",
        "- Propane: for kitchen, bar, and any registered flame effects — refill at Hell Station unless we qualify for delivery (see Section 9)",
        "- Estimated gallons per fuel type: complete with Infrastructure team — base on prior year burn rate plus contingency",
        "",
        "REQUIRED DIAGRAMS (upload with application):",
        "- Base camp diagram showing fuel storage location(s) relative to streets, RVs, structures, and ignition sources",
        "- Generator placement diagram showing 20 ft road access and 3 ft clearance on all sides",
        "- Propane storage location (must be ≥20 ft from any liquid fuel storage)",
        "",
        "Note: You will NOT be charged for estimated fuel — only for fuel actually dispensed. Estimate generously to ensure allocation, then pay daily on auto-pay for actual usage.",
    ]),
    (1, "6. Payment Account Setup (Gravity)", [
        "BMorg processes payments through Gravity Payments. Setup happens after registration approval (after July 1).",
        "",
        "1. Watch the Treasurer inbox after July 1 for the PETROL/Gravity activation email",
        "2. Click the activation link and create the PETROL Payment Account",
        "3. Save the camp credit card and designate it for AUTO-PAY (required — declined charges delay fueling on playa)",
        "4. Pay the account setup fee (one-time, covers saving the card on file)",
        "5. Confirm the Payment Contact name on file (cardholder must be on playa for activation, or notify PETROL in advance)",
        "6. Keep the card active until BMorg notifies us the account is paid in full — charges continue post-event during reconciliation",
        "7. Save Customer Portal login credentials to the camp shared drive (Treasurer + Infrastructure Director both need access)",
    ]),
    (1, "7. On-Playa Account Activation & Asset Registration", [
        "Nothing happens on playa until the account is activated in person. Deliveries do not start until the day AFTER assets are tagged.",
        "",
        "ACTIVATION (Infrastructure Director, day of arrival):",
        "1. Go to Hell Station Office (9:30 & halfway to fenceline) during office hours (10:00 am – 5:00 pm in 2025)",
        "2. Bring a government-issued ID — physical card only; photo on cup is NOT accepted",
        "3. Identify yourself as one of the designated account users (Project Lead, Fuel Contact, or Payment Contact)",
        "4. Confirm payment method is still on file and auto-pay is active",
        "5. Note the field registration ETA — agents are busiest Thursday pre-event through Monday of event",
        "",
        "FIELD REGISTRATION (24+ hours before first delivery):",
        "1. Field Registration Agents visit camp to inspect approved assets",
        "2. Camp must be set up to match the approved layout — fuel storage, generator, propane area all in their approved positions",
        "3. Storage must meet Section 8 requirements before agents arrive — non-compliance means no fuel",
        "4. Agents tag the approved assets and add them to the delivery route",
        "5. First fuel delivery comes the day AFTER tagging — routes are set for the day, no same-day adds",
    ]),
    (1, "8. Camp Storage & Safety Requirements", [
        "These are non-negotiable. If any of these fail at Field Registration, we will NOT be fueled.",
        "",
        "STORAGE QUANTITY LIMITS (separate from fuel-in-use in generators/RVs):",
        "- Liquid fuel (diesel): max 110 gal stored OR two 55-gal drums",
        "- Propane: max 500 lb (≈120 gal) stored",
        "- Liquid and propane limits are NOT cumulative — they are independent",
        "",
        "DISTANCES (rigid):",
        "- 10 ft minimum from any stored fuel to combustibles (tents, shade) or ignition sources (RVs, cars, trailers)",
        "- 20 ft minimum between liquid fuel storage and propane storage",
        "- 50 ft minimum between separate fuel storage areas (ours and neighbors')",
        "- 20 ft fire lane kept clear (camps larger than 100' x 100' only)",
        "- 10 ft buffer from city streets to prevent vehicle collision; mark with caution tape",
        "",
        "GENERATOR PLACEMENT (for in-camp delivery):",
        "- Within 20 ft of a lettered/numbered street with clear, unobstructed access (truck hose max reach)",
        "- 3 ft clear working space on all sides (also serves as ventilation)",
        "- Minimum 30-gallon delivery capacity per fueling",
        "- Minimum 30 hours runtime between deliveries",
        "- Maximum two generators per camp (or two aggregation locations of same-sized generators)",
        "- Fuel trucks do NOT back up or turn around in camp — provide a straight pull-through or 20 ft turn allowances",
        "- Keep bikes/vehicles clear of generator access during fueling hours or we get skipped",
        "",
        "SECONDARY CONTAINMENT (required for all liquid fuel storage):",
        "- Capacity = 110% of the largest single container",
        "- Examples: two 55-gal drums → 60.5 gal containment; five 5-gal cans → 5.5 gal containment",
        "- Must be impermeable to the fuel stored, free of cracks, and allow liquid removal",
        "- Truck beds do NOT count as secondary containment",
        "- Propane cylinders do NOT require secondary containment",
        "",
        "REQUIRED ON-SITE:",
        "- 'NO SMOKING — FLAMMABLE' signage visible from all four directions",
        "- Minimum one hand-held 40-B-rated fire extinguisher per fuel storage area, within 8–10 ft (pole-mounted is best)",
        "- Spill kit with shovel and sealable container for contaminated playa",
        "- Approved fuel containers, closed except when filling/dispensing/venting",
        "- For drums: own hand pump compatible with fuel type, fitting 2\" NPT bung port; nozzle must store ABOVE drum height to prevent siphon",
        "- Grounding rod for any gasoline drum (drums must be bonded together and to filling containers)",
        "",
        "SPILL RESPONSE:",
        "- Small spill (fits in a 5-gal bucket): shovel contaminated playa into spill kit container, pack out post-event",
        "- Larger spill: stop ignition sources, hold position, report to Black Rock Rangers, Earth Guardians, or ESD with location + contact person",
    ]),
    (1, "9. Diesel — Generator Delivery & Drum Program", [
        "Two delivery options for diesel. Decide on the application; defaults below.",
        "",
        "DIRECT GENERATOR DELIVERY (preferred for active generators):",
        "- Truck pulls up to the generator and fills directly during fueling routes",
        "- Available Monday pre-event through final Sunday",
        "- Deliveries Monday/Tuesday post-event require scheduling in advance with the Accounts team",
        "- No additional fees beyond the per-gallon dispensed price",
        "",
        "DRUM PROGRAM (55-gal drums, dispensed as 45 gal for heat expansion):",
        "- BMorg delivers a sealed drum to the camp fuel area; we dispense via our own hand pump",
        "- Drum delivery window: Wednesday pre-event through Friday of event",
        "- We can bring our own drums (subject to inspection) or use BMorg drums (free rental)",
        "- Fees per drum: $15 delivery + $20 secondary containment (if rented) + $9 safety kit (grounding/venting)",
        "- DO NOT use drums to dispose of any other substance — hazardous disposal fees will be charged",
        "- Pickup: starts Sunday Temple Burn day. Final hard deadline: drop at Hell Station by Monday 4:00 pm. Never leave drums unattended at strike.",
        "",
        "Note: BMorg supplies dyed diesel ($7.67/gal in 2024 — for reference only; 2026 prices announced in mid-August and updated daily) for off-road/generator use.",
    ]),
    (1, "10. Propane — Refueling Options", [
        "Propane refueling is location-restricted. Default option for NC is Hell Station pickup unless we qualify for delivery.",
        "",
        "HELL STATION PICKUP (default):",
        "- Bring CARB/EPA-certified cylinders to Hell Station during fueling hours",
        "- Transport in secondary containment (a bike trailer with a sealed bin works)",
        "- Propane lanes open ~one day after liquid fuel lanes (Aug 20 in 2025)",
        "",
        "IN-CAMP DELIVERY (only if eligible):",
        "- Available ONLY for cylinders 100 lb or larger, or tanks 25 gal or larger",
        "- Delivery locations restricted to Esplanade, 2:00 road, 10:00 road, outermost 9:00 street, some B Street keyhole frontages",
        "- Tanks must be within 50 ft of road with clear access",
        "- Only available for camps with registered flame effects",
        "",
        "PROPANE SAFETY:",
        "- All cylinders upright and secured against tipping",
        "- Cylinder requalification dates must be current; no excessive rust",
        "- Consumer cylinders 4–40 lb must have OPD (Overfill Protection Device) valves",
        "- Valve protection caps stay on except during use",
        "- Do NOT extinguish a leaking propane fire unless the leak can be stopped — call Rangers + ESD response team",
    ]),
    (1, "11. Pricing & Daily Billing", [
        "Prices are market-based and announced in mid-August; final daily rates posted at Hell Station and on the Customer Portal.",
        "",
        "2024 prices (reference only — 2026 will differ):",
        "- Gasoline 87: $8.07 / gal",
        "- Clear diesel: $7.91 / gal",
        "- Dyed diesel (generators): $7.67 / gal",
        "- Propane: $4.80 / gal",
        "",
        "BILLING:",
        "- Auto-pay charges the saved card daily during the event",
        "- Only dispensed gallons are charged (estimates are not billed)",
        "- Declined charges pause future fueling until resolved at Hell Station",
        "- Post-event charges continue during BMorg reconciliation — Treasurer keeps card active until 'paid in full' notification",
        "- Treasurer reconciles BMorg charges against camp accounting (see Tr1 once published)",
    ]),
    (1, "12. Record Keeping", [
        "Retain all records for a minimum of 5 years per nonprofit compliance requirements.",
        "",
        "- BMorg application confirmation (PDF) — pre-event",
        "- Approval letter from PETROL — pre-event",
        "- Gravity Payment Account credentials (stored in shared drive password vault, not the file tree)",
        "- Daily billing statements from Gravity / Customer Portal",
        "- Final post-event paid-in-full notification from BMorg",
        "- Reconciled fuel cost summary by fuel type and total gallons (for fiscal year close — see Tr2)",
        "- Any incident reports (spills, layout revisions, missed deliveries)",
        "",
        "File location: [Shared Drive > Treasurer > BMorg > PETROL > YYYY]",
    ]),
    (1, "13. Contacts", [
        "- BMorg PETROL Department: petrol@burningman.org",
        "- Hell Station (on-playa): 9:30 & halfway to fenceline; office 10:00 am – 5:00 pm",
        "- Treasurer (application + payment owner): Isabel Hoy — izhoy@yahoo.com",
        "- Infrastructure Director (on-playa fuel ops owner): [Name TBD] — [email]",
        "- President (oversight): Reece Dassinger — reece@nipplecrime.org",
        "- VP / SOP Owner: Chris Reddin — creddin1@hotmail.com",
    ]),
    (1, "14. Revision History", [
        "- v1.0 | 2026-05-13 | Initial draft | Chris Reddin",
    ]),
]

create_sop(
    output_path="Standard Operating Procedures/Tr5 Hell Station Petrol Application.docx",
    sop_number="Tr5",
    sop_title="Hell Station Petrol Application (BMorg)",
    department="Treasurer",
    version="1.0",
    effective_date="2026-05-13",
    last_updated="2026-05-13",
    sections=sections_tr5,
)
